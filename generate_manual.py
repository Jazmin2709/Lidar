
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ─── Helpers ─────────────────────────────────────────────────────────────────
def shade_paragraph(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def add_hr(doc, color_hex='2563EB'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading(doc, text, level=1, color_hex='1E3A5F'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else (13 if level == 2 else 11))
    r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.name = 'Calibri'
    return p

def add_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p

def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(1)
        shade_paragraph(p, 'EFF6FF')
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

def add_note(doc, text, label='📝 Nota', fill='DBEAFE', text_color='1E3A8A', label_color='1E40AF'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.5)
    p.paragraph_format.right_indent  = Cm(0.5)
    p.paragraph_format.space_before  = Pt(6)
    p.paragraph_format.space_after   = Pt(6)
    shade_paragraph(p, fill)
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = 'Calibri'
    lc = tuple(int(label_color[i:i+2], 16) for i in (0, 2, 4))
    r1.font.color.rgb = RGBColor(*lc)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = 'Calibri'
    tc = tuple(int(text_color[i:i+2], 16) for i in (0, 2, 4))
    r2.font.color.rgb = RGBColor(*tc)
    return p

def add_warning(doc, text):
    return add_note(doc, text, label='⚠️ Importante', fill='FEF3C7',
                    text_color='92400E', label_color='B45309')

def cell_shade(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_env_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell_shade(hdr_cells[i], '1E3A8A')
    for ri, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        fill = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            for para in row_cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            cell_shade(row_cells[i], fill)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
banner = doc.add_paragraph()
banner.paragraph_format.space_before = Pt(0)
banner.paragraph_format.space_after  = Pt(0)
shade_paragraph(banner, '1E3A8A')
banner.add_run('  ').font.size = Pt(30)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(30)
title_p.paragraph_format.space_after  = Pt(6)
t1 = title_p.add_run('MANUAL DE DESPLIEGUE')
t1.bold = True; t1.font.size = Pt(26); t1.font.name = 'Calibri'
t1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after  = Pt(30)
t2 = sub_p.add_run('SISTEMA LIDAR')
t2.bold = True; t2.font.size = Pt(18); t2.font.name = 'Calibri'
t2.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

add_hr(doc)

doc.add_paragraph()
ap = doc.add_paragraph()
ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
al = ap.add_run('Autor')
al.bold = True; al.font.size = Pt(13); al.font.name = 'Calibri'
al.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = auth_p.add_run('[Nombre del autor]')
ar.font.size = Pt(12); ar.font.name = 'Calibri'

doc.add_paragraph()
dept_p = doc.add_paragraph()
dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp = dept_p.add_run('Facultad de Ingeniería')
dp.font.size = Pt(12); dp.font.name = 'Calibri'
dp.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

year_p = doc.add_paragraph()
year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
yp = year_p.add_run('2025')
yp.font.size = Pt(12); yp.font.name = 'Calibri'
yp.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
banner2 = doc.add_paragraph()
shade_paragraph(banner2, '1E3A8A')
banner2.add_run('  ').font.size = Pt(30)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLA DE CONTENIDO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Contenido', level=1)
add_hr(doc)

toc_items = [
    ('1.', 'Objetivo', '3'),
    ('2.', 'Requisitos previos', '3'),
    ('2.1', 'Cuentas y servicios necesarios', '3'),
    ('2.2', 'Herramientas locales', '4'),
    ('3.', 'Configuración de la base de datos – MySQL en AWS RDS', '4'),
    ('3.1', 'Crear instancia RDS', '4'),
    ('3.2', 'Configurar parámetros de seguridad', '5'),
    ('3.3', 'Ejecutar script SQL inicial', '5'),
    ('4.', 'Despliegue del Backend – Render (Web Service)', '5'),
    ('4.1', 'Preparar el repositorio del backend', '5'),
    ('4.2', 'Crear el servicio en Render', '6'),
    ('4.3', 'Variables de entorno del backend (.env)', '6'),
    ('5.', 'Despliegue del Frontend – Render (Static Site)', '7'),
    ('5.1', 'Preparar el repositorio del frontend', '7'),
    ('5.2', 'Variables de entorno del frontend (Vite)', '7'),
    ('5.3', 'Crear el sitio estático en Render', '9'),
    ('6.', 'Verificación del despliegue', '9'),
    ('7.', 'Solución de problemas comunes', '10'),
]

toc_table = doc.add_table(rows=len(toc_items), cols=2)
toc_table.style = 'Table Grid'
for ri, (num, label, page) in enumerate(toc_items):
    cells = toc_table.rows[ri].cells
    cells[0].text = f'{num}  {label}'
    cells[1].text = page
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fill = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
    for ci in range(2):
        cell_shade(cells[ci], fill)
        tc = cells[ci]._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
        tcPr.append(tcBorders)
        for para in cells[ci].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10.5)
                run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. OBJETIVO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Objetivo', level=1)
add_hr(doc)
add_body(doc,
    'El objetivo de este documento es describir de forma detallada el proceso de despliegue del '
    'Sistema LIDAR en infraestructura en la nube, abarcando la base de datos en Amazon Web Services '
    '(AWS RDS), el backend como servicio web en Render y el frontend como sitio estático también en '
    'Render. Este manual sirve como guía de referencia para el equipo de desarrollo y operaciones '
    'a fin de reproducir o actualizar el entorno productivo de manera confiable y ordenada.'
)

# ═══════════════════════════════════════════════════════════════════════════════
#  2. REQUISITOS PREVIOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Requisitos previos', level=1)
add_hr(doc)

add_heading(doc, '2.1  Cuentas y servicios necesarios', level=2)
add_body(doc, 'Para llevar a cabo el despliegue se requieren las siguientes cuentas activas:')
for item in [
    'Cuenta en AWS (https://aws.amazon.com) con permisos para crear instancias RDS.',
    'Cuenta en Render (https://render.com) – plan gratuito o de pago.',
    'Repositorio del proyecto en GitHub (https://github.com) con acceso de lectura/escritura.',
    'Cuenta en Cloudinary (https://cloudinary.com) para el almacenamiento de imágenes.',
    'Clave de API de Groq (https://console.groq.com) para el servicio de chatbot.',
]:
    add_bullet(doc, item)

add_heading(doc, '2.2  Herramientas locales', level=2)
add_body(doc, 'En la máquina del desarrollador se requiere:')
for item in [
    'Node.js v18 o superior (https://nodejs.org)',
    'npm v9 o superior (incluido con Node.js)',
    'Git (https://git-scm.com)',
    'Cliente MySQL (MySQL Workbench o línea de comandos)',
    'Acceso a Internet para conectarse a los servicios en la nube.',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. BASE DE DATOS – AWS RDS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. Configuración de la base de datos – MySQL en AWS RDS', level=1)
add_hr(doc)

add_heading(doc, '3.1  Crear instancia RDS', level=2)
add_body(doc, 'Acceder a la consola de AWS y navegar a RDS → Databases → Create database. Seguir los pasos indicados:')
for s in [
    'Seleccionar el motor: MySQL (versión 8.0).',
    'Elegir la plantilla: Free tier (para pruebas) o Production (para producción).',
    'Configurar el identificador de instancia, por ejemplo: lidar-db.',
    'Definir un nombre de usuario maestro (p. ej. admin) y una contraseña segura.',
    'En Connectivity → Public access: seleccionar "Yes" si se necesita acceso externo durante el desarrollo.',
    'Anotar el Endpoint generado por AWS (p. ej. lidar-db.xxxxxxxxx.us-east-1.rds.amazonaws.com).',
    'Hacer clic en Create database y esperar a que el estado cambie a "Available".',
]:
    add_bullet(doc, s)

add_heading(doc, '3.2  Configurar parámetros de seguridad', level=2)
for s in [
    'Ir a EC2 → Security Groups y seleccionar el grupo asociado a la instancia RDS.',
    'En Inbound rules, agregar una regla de tipo MySQL/Aurora (puerto 3306).',
    'Como origen, agregar la IP del desarrollador y los rangos de IP de Render.',
]:
    add_bullet(doc, s)

add_heading(doc, '3.3  Ejecutar el script SQL inicial', level=2)
add_body(doc, 'Conectarse al endpoint de RDS con MySQL Workbench o la línea de comandos y ejecutar el script lidarN.sql:')
add_code(doc, ['mysql -h <ENDPOINT_RDS> -u admin -p < lidarN.sql'])
add_note(doc,
    'El repositorio contiene tres versiones del script SQL (lidar.sql, lidar3.sql, lidarN.sql). '
    'Se debe utilizar lidarN.sql ya que corresponde a la versión más actualizada del esquema.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. BACKEND – RENDER WEB SERVICE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Despliegue del Backend – Render (Web Service)', level=1)
add_hr(doc)
add_body(doc,
    'El backend del Sistema LIDAR está desarrollado con Node.js y Express. Se despliega en Render '
    'como un Web Service conectado directamente al repositorio de GitHub.'
)

add_heading(doc, '4.1  Preparar el repositorio del backend', level=2)
add_body(doc, 'Verificar lo siguiente en el repositorio antes de crear el servicio:')
for item in [
    'El archivo package.json tiene definido el script start: "nodemon index.js" (o "node index.js" en producción).',
    'El puerto del servidor usa la variable de entorno PORT: const PORT = process.env.PORT || 3000;',
    'El archivo .gitignore excluye node_modules/ y el archivo .env (ya configurado en el proyecto).',
    'Todas las credenciales se cargan con dotenv mediante process.env desde el archivo .env.',
]:
    add_bullet(doc, item)

add_heading(doc, '4.2  Crear el servicio en Render', level=2)
for step in [
    'Iniciar sesión en https://render.com y hacer clic en New → Web Service.',
    'Seleccionar "Connect a repository" y otorgar acceso al repositorio del backend en GitHub.',
    'Configurar los parámetros del servicio:',
]:
    add_bullet(doc, step)

add_env_table(doc,
    ['Parámetro', 'Valor recomendado'],
    [
        ['Name', 'lidar-backend'],
        ['Region', 'Oregon (US West) u otra cercana'],
        ['Branch', 'main'],
        ['Root Directory', 'backend (si es monorepo)'],
        ['Runtime', 'Node'],
        ['Build Command', 'npm install'],
        ['Start Command', 'node index.js'],
        ['Instance Type', 'Free (o Starter para producción)'],
    ]
)
for step in [
    'Hacer clic en Create Web Service. Render instalará dependencias y desplegará el servidor.',
    'Al finalizar, Render asigna una URL pública: https://lidar-backend.onrender.com',
    'Anotar esta URL: se usará como valor de VITE_API_URL en el frontend.',
]:
    add_bullet(doc, step)

add_heading(doc, '4.3  Variables de entorno del backend (.env)', level=2)
add_body(doc,
    'El backend usa dotenv para cargar variables de entorno. En desarrollo, se utiliza el archivo '
    'backend/.env (excluido del repositorio por .gitignore). En producción (Render), estas variables '
    'se configuran en el panel: Dashboard → Servicio → Environment → Add Environment Variable.'
)
add_note(doc,
    'El archivo .env ha sido añadido al .gitignore del backend. Nunca subas este archivo al repositorio '
    'ya que contiene credenciales sensibles. Cada desarrollador o servidor debe tener su propio .env.'
)

add_body(doc, 'Variables requeridas en el backend:', bold=True)
add_code(doc, [
    '# ── Base de datos (AWS RDS) ──────────────────────────────',
    'DB_HOST=lidar-db.xxxxxxxxx.us-east-1.rds.amazonaws.com',
    'DB_PORT=3306',
    'DB_NAME=lidar',
    'DB_USER=admin',
    'DB_PASS=tu_contrasena_segura',
    '',
    '# ── JWT ───────────────────────────────────────────────────',
    'JWT_SECRET=clave_muy_segura_123',
    '',
    '# ── Cloudinary (almacenamiento de imágenes) ───────────────',
    'CLOUDINARY_CLOUD_NAME=mi_cloud',
    'CLOUDINARY_API_KEY=123456789012345',
    'CLOUDINARY_API_SECRET=abc123xyz',
    '',
    '# ── Groq (chatbot IA) ────────────────────────────────────',
    'GROQ_API_KEY=gsk_xxxxxxxxxxxx',
    '',
    '# ── Correo (nodemailer) ──────────────────────────────────',
    'EMAIL_USER=noreply@dominio.com',
    'EMAIL_PASS=contrasena_del_correo',
    '',
    '# ── Puerto (Render lo asigna automáticamente) ────────────',
    '# PORT=3000   ← NO definir manualmente en Render',
])
add_warning(doc,
    'Render inyecta automáticamente la variable PORT. No es necesario definirla manualmente '
    'en el panel de Render, solo asegúrate de que el código la use: process.env.PORT || 3000.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. FRONTEND – RENDER STATIC SITE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Despliegue del Frontend – Render (Static Site)', level=1)
add_hr(doc)
add_body(doc,
    'El frontend del Sistema LIDAR está desarrollado con React y Vite. Se despliega en Render como '
    'un Static Site. El código fuente ha sido refactorizado para eliminar URLs hardcodeadas; '
    'todas las llamadas al backend se realizan a través de una variable de entorno dinámica.'
)

add_heading(doc, '5.1  Preparar el repositorio del frontend', level=2)
for item in [
    'Verificar que el archivo vite.config.js esté configurado correctamente con el plugin de React.',
    'Confirmar que node_modules/ y .env estén en el .gitignore del frontend.',
    'Hacer commit y push de todos los cambios al repositorio.',
]:
    add_bullet(doc, item)

add_heading(doc, '5.2  Variables de entorno del frontend (Vite)', level=2)
add_body(doc,
    'A diferencia del backend (Node.js), el frontend usa Vite y no tiene acceso al objeto '
    'process.env de Node. Vite tiene su propio sistema de variables de entorno basado en '
    'import.meta.env, con la restricción de que solo las variables que comiencen por '
    'VITE_ son expuestas al navegador por razones de seguridad.'
)
add_warning(doc,
    'En proyectos con Vite NO se usa process.env para acceder a variables en el código del cliente. '
    'El uso de process.env en el frontend resultará en el error "process is not defined" en el navegador, '
    'ya que process es un objeto exclusivo de Node.js.'
)

add_body(doc, 'Comparativa: Backend vs Frontend', bold=True)
add_env_table(doc,
    ['Entorno', 'Herramienta', 'Sintaxis de acceso', 'Prefijo requerido'],
    [
        ['Backend (Node.js)', 'dotenv (.env)', 'process.env.VARIABLE', 'Ninguno'],
        ['Frontend (Vite)', 'Vite env (.env)', 'import.meta.env.VITE_VARIABLE', 'VITE_ obligatorio'],
    ]
)

add_body(doc, 'Archivo frontend/.env (desarrollo local):', bold=True)
add_code(doc, [
    '# URL base del backend (apunta al servidor local durante el desarrollo)',
    'VITE_API_URL=http://localhost:3000/api',
])
add_note(doc,
    'El archivo frontend/.env está excluido del repositorio por .gitignore. '
    'En Render, esta variable se configura directamente en el panel del Static Site. '
    'El código del proyecto está preparado para usar un valor de respaldo si la variable no está definida: '
    'import.meta.env.VITE_API_URL || "http://localhost:3000/api"'
)

add_body(doc, 'Uso en el código fuente del frontend:', bold=True)
add_body(doc, 'Todos los archivos que realizan peticiones al backend utilizan el siguiente patrón:')
add_code(doc, [
    '// Patrón estándar aplicado en todos los archivos del frontend',
    'const API_URL = import.meta.env.VITE_API_URL || "http://localhost:3000/api";',
    '',
    '// Ejemplo de uso en BuddyFormulario.jsx',
    'const API_URL = import.meta.env.VITE_API_URL || "http://localhost:3000/api";',
    'const BUDDY_API_URL = `${API_URL}/buddy`;',
    '',
    '// Ejemplo de uso en Reportes.jsx',
    'const API_URL = import.meta.env.VITE_API_URL || "http://localhost:3000/api";',
    'const API_BASE_URL = `${API_URL}/buddy/BuddyPartner`;',
    'const API_IMAGE_UPLOAD = `${API_URL}/imagenes/subir`;',
])

add_body(doc, 'Archivos del frontend actualizados con este patrón:', bold=True)
for f in [
    'src/paginas/BuddyFormulario.jsx',
    'src/paginas/Reportes.jsx',
    'src/paginas/RecuperarContraseña.jsx',
    'src/paginas/PagInicio.jsx',
    'src/paginas/IndexEmpleado.jsx',
    'src/paginas/EnviarCorreo.jsx',
    'src/paginas/Empleados.jsx',
    'src/paginas/Dashboard.jsx',
    'src/componentes/RoutePublic.jsx',
    'src/componentes/RoutePrivate.jsx',
    'src/componentes/ChatBot.jsx',
]:
    add_bullet(doc, f)

add_heading(doc, '5.3  Crear el sitio estático en Render', level=2)
for step in [
    'Iniciar sesión en https://render.com y hacer clic en New → Static Site.',
    'Conectar el repositorio de GitHub del frontend.',
    'Configurar los parámetros:',
]:
    add_bullet(doc, step)

add_env_table(doc,
    ['Parámetro', 'Valor recomendado'],
    [
        ['Name', 'lidar-frontend'],
        ['Branch', 'main'],
        ['Root Directory', 'frontend (si es monorepo)'],
        ['Build Command', 'npm install && npm run build'],
        ['Publish Directory', 'dist'],
    ]
)
for step in [
    'En Environment, agregar la variable VITE_API_URL con la URL pública del backend '
    '(p. ej. https://lidar-backend.onrender.com/api).',
    'Importante: añadirla ANTES del primer build, ya que Vite la embebe en el bundle estático durante la compilación.',
    'En Redirects/Rewrites, agregar una regla: Source /* → Destination /index.html → Action Rewrite. '
    'Esto es necesario para que React Router maneje las rutas del lado del cliente.',
    'Hacer clic en Create Static Site. Render ejecutará npm run build y publicará el contenido de dist/.',
    'Al finalizar, se asigna una URL pública: https://lidar-frontend.onrender.com',
]:
    add_bullet(doc, step)

add_warning(doc,
    'A diferencia de las variables del backend, VITE_API_URL es procesada en tiempo de BUILD, '
    'no en tiempo de ejecución. Si cambias su valor en el panel de Render, debes hacer un nuevo '
    'deploy (re-build) para que el cambio surta efecto en el frontend.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. VERIFICACIÓN
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Verificación del despliegue', level=1)
add_hr(doc)
add_body(doc, 'Realizar las siguientes comprobaciones para confirmar que todos los componentes están operativos:')

for title, desc in [
    ('Base de datos (RDS)',
     'Conectarse con MySQL Workbench al endpoint de RDS y confirmar que las tablas del esquema existen.'),
    ('Backend (Render Web Service)',
     'Acceder a la URL del backend desde el navegador. Debe responder: "API de Lidar funcionando 🚀"'),
    ('Autenticación',
     'Realizar POST a /api/auth/login con credenciales de prueba y verificar que retorna un token JWT.'),
    ('Frontend (Render Static Site)',
     'Acceder a la URL del sitio estático y verificar que la pantalla de inicio carga sin errores en consola.'),
    ('Variable VITE_API_URL',
     'Desde las DevTools del navegador, verificar en la pestaña Network que las peticiones se dirigen a la URL del backend en Render (no a localhost).'),
    ('Flujo completo',
     'Iniciar sesión desde el frontend y navegar por las secciones para confirmar la comunicación frontend → backend → base de datos.'),
    ('Carga de imágenes',
     'Subir una imagen de prueba en el formulario Buddy Partner y verificar que se almacena en Cloudinary.'),
    ('Chatbot',
     'Interactuar con el chatbot para verificar la conectividad con Groq.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(f'{title}: ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Calibri'
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.name = 'Calibri'

doc.add_paragraph()
add_warning(doc,
    'En el plan gratuito de Render, los Web Services se suspenden tras 15 minutos sin tráfico. '
    'La primera petición puede tardar hasta 50 segundos mientras el servicio se reactiva. '
    'Se recomienda el plan Starter para entornos de producción. Los Static Sites no tienen esta limitación.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. SOLUCIÓN DE PROBLEMAS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. Solución de problemas comunes', level=1)
add_hr(doc)

problems = [
    ('VITE_API_URL apunta a localhost en producción',
     'Verificar que la variable VITE_API_URL esté definida en el panel de Render del Static Site ANTES del build. '
     'Si se añadió después del primer deploy, hacer un nuevo deploy manual desde el panel.',
     'Variable añadida después del build o con nombre incorrecto (sin prefijo VITE_).'),

    ('process.env no definido en el frontend',
     'En proyectos Vite, NO se usa process.env. Reemplazar con import.meta.env.VITE_NOMBRE_VARIABLE. '
     'Si el nombre de la variable no empieza con VITE_, Vite la ignorará y retornará undefined.',
     'Se intentó usar la sintaxis de Node.js (process.env) en código del cliente Vite.'),

    ('Error de conexión a la base de datos',
     'Verificar que el Security Group de RDS permita tráfico en el puerto 3306 desde las IPs de Render. '
     'Confirmar que las variables DB_HOST, DB_USER, DB_PASS y DB_NAME estén correctas en el backend.',
     'Security Group bloqueando la conexión o variables de entorno incorrectas.'),

    ('Build del frontend falla en Render',
     'Revisar los logs de build en Render. Verificar que la variable VITE_API_URL esté definida. '
     'Confirmar que npm run build funciona localmente antes de hacer push.',
     'Dependencias faltantes o variables de entorno no configuradas.'),

    ('CORS: peticiones bloqueadas desde el frontend',
     'En index.js del backend, actualizar origin de CORS para incluir la URL del frontend en Render: '
     'origin: "https://lidar-frontend.onrender.com". En desarrollo se puede usar origin: "*".',
     'La URL del frontend cambió o CORS está configurado con un origen diferente.'),

    ('Página en blanco en el frontend (React Router)',
     'Agregar la regla de rewrite en Render Static Site: /* → /index.html. '
     'Esto es necesario para que React Router gestione las rutas del lado del cliente.',
     'Falta la regla de rewrite en la configuración del Static Site.'),

    ('Imágenes no se suben a Cloudinary',
     'Verificar que las variables CLOUDINARY_CLOUD_NAME, CLOUDINARY_API_KEY y CLOUDINARY_API_SECRET '
     'estén definidas correctamente en el backend de Render.',
     'Credenciales de Cloudinary incorrectas o no definidas en el entorno de producción.'),

    ('El chatbot no responde',
     'Verificar que la variable GROQ_API_KEY esté definida en el panel de Render del backend '
     'y que la clave sea válida en la consola de Groq.',
     'Clave de API expirada, inválida o no definida en el entorno.'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, h in enumerate(['Problema', 'Solución', 'Causa probable']):
    hdr[i].text = h
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell_shade(hdr[i], '1E3A8A')

for ri, (prob, sol, cause) in enumerate(problems):
    row_cells = table.add_row().cells
    row_cells[0].text = prob
    row_cells[1].text = sol
    row_cells[2].text = cause
    fill = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
    for i in range(3):
        for para in row_cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5); run.font.name = 'Calibri'
        cell_shade(row_cells[i], fill)

# ─── Save ─────────────────────────────────────────────────────────────────────
output_path = r'c:\xampp\htdocs\Lidar\Manual_Despliegue_LIDAR.docx'
doc.save(output_path)
print(f'✅ Manual actualizado en: {output_path}')
